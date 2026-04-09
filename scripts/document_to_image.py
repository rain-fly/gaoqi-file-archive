#!/usr/bin/env python3
"""
文档转图片模块 - 将 PDF/Word 文档转为图片

使用 PyMuPDF (fitz) 处理 PDF
使用 python-docx + matplotlib 处理 Word (.docx)

注意：.doc 格式不支持，需要先转换为 .docx
"""
import os
import tempfile
from pathlib import Path
from typing import Optional, Tuple

# 尝试导入依赖
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    print("警告: PyMuPDF 未安装，PDF 转换将不可用")

try:
    from docx import Document
    import matplotlib.pyplot as plt
    import matplotlib.patches as patches
    from PIL import Image
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("警告: python-docx 或 matplotlib 未安装，Word 转换将不可用")


# 支持的文档扩展名
SUPPORTED_EXTS = {'.pdf', '.docx'}


def pdf_first_page_to_image(pdf_path: str,
                             output_path: Optional[str] = None,
                             dpi: int = 150) -> Optional[str]:
    """
    将 PDF 第一页转换为图片

    Args:
        pdf_path: PDF 文件路径
        output_path: 输出图片路径（可选，默认生成临时文件）
        dpi: 图片分辨率

    Returns:
        输出图片路径，失败返回 None
    """
    if not HAS_PYMUPDF:
        print("错误: PyMuPDF 未安装")
        return None

    if not os.path.exists(pdf_path):
        print(f"错误: 文件不存在 {pdf_path}")
        return None

    try:
        # 打开 PDF
        doc = fitz.open(pdf_path)
        if doc.page_count == 0:
            print("错误: PDF 页面数为0")
            return None

        # 获取第一页
        page = doc[0]

        # 计算缩放因子（基于 DPI）
        zoom = dpi / 72  # PDF 标准是 72 DPI
        mat = fitz.Matrix(zoom, zoom)

        # 渲染页面为图片
        pix = page.get_pixmap(matrix=mat)

        # 确定输出路径
        if output_path is None:
            fd, output_path = tempfile.mkstemp(suffix='.png')
            os.close(fd)
        else:
            # 确保输出目录存在
            os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)

        # 保存图片
        pix.save(output_path)
        doc.close()

        return output_path

    except Exception as e:
        print(f"PDF 转换失败: {e}")
        return None


def docx_first_page_to_image(docx_path: str,
                              output_path: Optional[str] = None,
                              width_px: int = 800) -> Optional[str]:
    """
    将 Word 文档第一页转换为图片

    Args:
        docx_path: Word 文件路径
        output_path: 输出图片路径（可选，默认生成临时文件）
        width_px: 输出图片宽度（像素）

    Returns:
        输出图片路径，失败返回 None
    """
    if not HAS_DOCX:
        print("错误: python-docx 或 matplotlib 未安装")
        return None

    if not os.path.exists(docx_path):
        print(f"错误: 文件不存在 {docx_path}")
        return None

    try:
        # 打开文档
        doc = Document(docx_path)

        # 创建图形
        fig, ax = plt.subplots(figsize=(8.27, 11.69))  # A4 大小
        ax.set_xlim(0, 100)
        ax.set_ylim(0, 130)
        ax.axis('off')

        # 标题
        title = doc.core_properties.title or Path(docx_path).stem
        ax.text(50, 125, title, ha='center', va='top',
                fontsize=14, fontweight='bold')

        # 收集段落文本（只取第一页内容，简化处理）
        y_position = 115
        line_height = 3
        paragraphs_collected = []
        total_chars = 0
        max_chars = 1500  # 限制内容长度

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 简化处理：直接添加文本，不考虑复杂格式
            paragraphs_collected.append(text)
            total_chars += len(text)

            if total_chars > max_chars:
                break

        # 绘制文本
        text_content = '\n'.join(paragraphs_collected)
        ax.text(5, y_position, text_content, ha='left', va='top',
                fontsize=8, wrap=True, family='monospace')

        # 收集表格（如果有）
        tables_text = []
        for i, table in enumerate(doc.tables[:3]):  # 最多3个表格
            table_text = f"[表格 {i+1}]"
            for row in table.rows[:5]:  # 每个表格最多5行
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    table_text += '\n' + row_text
            tables_text.append(table_text)
            if i >= 2:
                break

        # 绘制表格
        if tables_text:
            table_y = y_position - len(paragraphs_collected) * 3 - 5
            for table_text in tables_text:
                ax.text(5, table_y, table_text, ha='left', va='top',
                        fontsize=7, family='monospace',
                        bbox=dict(boxstyle='round', facecolor='lightgray', alpha=0.3))
                table_y -= 20

        # 调整布局
        plt.tight_layout()

        # 确定输出路径
        if output_path is None:
            fd, output_path = tempfile.mkstemp(suffix='.png')
            os.close(fd)
        else:
            # 确保输出目录存在
            os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)

        # 保存图片
        fig.savefig(output_path, dpi=150, bbox_inches='tight',
                    facecolor='white', edgecolor='none')
        plt.close(fig)

        return output_path

    except Exception as e:
        print(f"Word 转换失败: {e}")
        return None


def convert_first_page_to_image(file_path: str,
                                output_path: Optional[str] = None,
                                dpi: int = 150) -> Optional[str]:
    """
    转换文档第一页为图片（自动判断类型）

    Args:
        file_path: 文档路径
        output_path: 输出图片路径（可选）
        dpi: PDF 分辨率

    Returns:
        输出图片路径，失败返回 None
    """
    ext = Path(file_path).suffix.lower()

    if ext == '.pdf':
        return pdf_first_page_to_image(file_path, output_path, dpi)
    elif ext == '.docx':
        return docx_first_page_to_image(file_path, output_path)
    elif ext == '.doc':
        print("错误: .doc 格式不支持，请先转换为 .docx")
        return None
    else:
        print(f"错误: 不支持的格式 {ext}")
        return None


def is_document_file(file_path: str) -> bool:
    """判断是否为文档文件"""
    return Path(file_path).suffix.lower() in SUPPORTED_EXTS


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("用法: python document_to_image.py <文档路径> [输出图片路径]")
        sys.exit(1)

    file_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None

    print(f"输入: {file_path}")
    result = convert_first_page_to_image(file_path, output_path)

    if result:
        print(f"输出: {result}")
    else:
        print("转换失败")
